import docx
import os

processed_path = "/home/erick-fcs/Descargas/Practicas_preprofesionales/processed/pregrado/25.  INFORME SG DERECHO 2023.docx"
if os.path.exists(processed_path):
    doc = docx.Document(processed_path)
    print("--- Processed Table 0 content ---")
    for r in doc.tables[0].rows:
        print(" | ".join(c.text.strip().replace("\n", " ") for c in r.cells))
        
    print("\n--- Search population count paragraph ---")
    for idx, p in enumerate(doc.paragraphs):
        if "seguimiento a graduados cohorte" in p.text.lower():
            print(f"P{idx}: {p.text}")
else:
    print("File not found:", processed_path)
