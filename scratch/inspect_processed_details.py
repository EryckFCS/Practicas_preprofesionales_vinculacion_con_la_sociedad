import docx
import os

doc_path = "/home/erick-fcs/Descargas/Practicas_preprofesionales/processed/pregrado/26.  INFORME SG ECONOMIA 2023.docx"
doc = docx.Document(doc_path)
print("=== PROCESSED DOCUMENT PARAGRAPHS 110-145 ===")
for i in range(110, 145):
    if i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        print(f"[{i}]: Style={repr(p.style.name)} Text={repr(p.text)}")
