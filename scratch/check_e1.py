import docx
import os

doc_path = "/home/erick-fcs/Descargas/Practicas_preprofesionales/INFORMES 24-46 ERICK/25.  INFORME SG DERECHO 2023.docx"
if not os.path.exists(doc_path):
    print("File not found:", doc_path)
    exit(1)

doc = docx.Document(doc_path)
print(f"Total paragraphs: {len(doc.paragraphs)}")

for idx, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if "3.5" in text or "E1" in text or "capacitación" in text.lower() or "conocimientos en …" in text or "…." in text:
        print(f"[{idx}] (Length: {len(text)}): {text}")
