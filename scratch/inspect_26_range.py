import docx
import os

doc_path = "/home/erick-fcs/Descargas/Practicas_preprofesionales/INFORMES 24-46 ERICK/26.  INFORME SG ECONOMIA 2023.docx"
doc = docx.Document(doc_path)
for i in range(120, 140):
    if i < len(doc.paragraphs):
        print(f"[{i}]: {repr(doc.paragraphs[i].text)}")
