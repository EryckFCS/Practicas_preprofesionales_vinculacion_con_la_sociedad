import docx

doc = docx.Document("/home/erick-fcs/Descargas/Practicas_preprofesionales/processed/25.  INFORME SG DERECHO 2023.docx")
for idx, p in enumerate(doc.paragraphs):
    if "sugirieron las siguientes temáticas" in p.text:
        print(f"Found intro at P{idx}")
        for i in range(idx, min(idx + 35, len(doc.paragraphs))):
            print(f"P{i}: {doc.paragraphs[i].text}")
        break
