import docx
import os

doc_path = "/home/erick-fcs/Descargas/Practicas_preprofesionales/INFORMES 24-46 ERICK/26.  INFORME SG ECONOMIA 2023.docx"
if not os.path.exists(doc_path):
    print("File not found:", doc_path)
    exit(1)

doc = docx.Document(doc_path)
print(f"Total paragraphs: {len(doc.paragraphs)}")

for idx, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if len(text) > 0:
        # Check for matching headers or text relating to "enseñanza-aprendizaje" or "perfil de egreso" or "4.5"
        if "4.5" in text or "enseñanza" in text.lower() or "perfil de egreso" in text.lower() or "satisfacción" in text.lower() or "egresado" in text.lower():
            print(f"[{idx}]: {text}")
