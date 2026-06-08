import docx
import os

word_path = "/home/erick-fcs/Descargas/Practicas_preprofesionales/INFORMES 24-46 ERICK/25.  INFORME SG DERECHO 2023.docx"
if os.path.exists(word_path):
    doc = docx.Document(word_path)
    print("Num tables:", len(doc.tables))
    print("--- Table 0 content ---")
    for r in doc.tables[0].rows:
        print(" | ".join(c.text.strip().replace("\n", " ") for c in r.cells))
        
    print("\n--- Search 89 in all paragraphs ---")
    for idx, p in enumerate(doc.paragraphs):
        if "89" in p.text:
            print(f"P{idx}: {p.text}")
            
    print("\n--- Search 89 in all tables ---")
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            row_str = " | ".join(c.text.strip().replace("\n", " ") for c in row.cells)
            if "89" in row_str:
                print(f"Table {t_idx} Row {r_idx}: {row_str}")
else:
    print("File not found:", word_path)
