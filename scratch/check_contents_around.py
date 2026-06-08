import docx
import os

for report_id in range(24, 35):
    career = {
        24: "24.  INFORME SG CONTAB. Y AUDIT 2023.docx",
        25: "25.  INFORME SG DERECHO 2023.docx",
        26: "26.  INFORME SG ECONOMIA 2023.docx",
        27: "27.  INFORME SG ENFERMERIA 2023.docx",
        28: "28.  INFORME SG LAB CLINICO 2023.docx",
        29: "29.  INFORME SG MEDICINA HUMANA 2023.docx",
        30: "30.  INFORME SG ODONTOLOGIA 2023.docx",
        31: "31.  INFORME SG PSICOL. CLINICA 2023.docx",
        32: "32.  INFORME SG ADM. EMPRESAS 2023.docx",
        33: "33.  INFORME SG CONTABILIDAD Y AUDIT. 2023.docx",
        34: "34.  INFORME SG DERECHO 2023.docx"
    }[report_id]
    doc_path = os.path.join("/home/erick-fcs/Descargas/Practicas_preprofesionales/INFORMES 24-46 ERICK", career)
    if os.path.exists(doc_path):
        doc = docx.Document(doc_path)
        print(f"\n================= REPORT {report_id} =================")
        # Find index of EA and PE
        ea_idx = -1
        pe_idx = -1
        for idx, p in enumerate(doc.paragraphs):
            t = p.text.strip()
            if t == "Percepción sobre el proceso de enseñanza-aprendizaje":
                ea_idx = idx
            elif t == "Percepción sobre el perfil de egreso":
                pe_idx = idx
        if ea_idx != -1 and pe_idx != -1:
            print(f"EA index: {ea_idx}, PE index: {pe_idx}")
            # Print paragraphs between ea_idx and pe_idx
            print(f"--- Text between EA and PE (should be EA content) ---")
            for i in range(ea_idx + 1, pe_idx):
                txt = doc.paragraphs[i].text.strip()
                if txt:
                    print(f"[{i}]: {txt}")
            # Print paragraphs after pe_idx up to next heading or 5 paragraphs
            print(f"--- Text after PE (should be PE content) ---")
            count = 0
            for i in range(pe_idx + 1, pe_idx + 6):
                if i < len(doc.paragraphs):
                    txt = doc.paragraphs[i].text.strip()
                    if txt:
                        print(f"[{i}]: {txt}")
                        if "calidad de" in txt.lower() or "servicios institucionales" in txt.lower():
                            break
    else:
        print(f"Report {report_id}: file not found")
