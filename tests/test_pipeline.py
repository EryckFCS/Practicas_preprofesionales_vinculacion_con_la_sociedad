import unittest
import os
import docx
from src.config import PARQUET_CACHE_DIR
from src.ingest import ingest_career_data
from src.extractor import extract_all_data
from src.updater import update_report

class TestPipeline(unittest.TestCase):
    
    @classmethod
    def setUpClass(cls):
        # Run ingestion once to prepare the Parquet cache for tests
        ingest_career_data(35, force=True)

    def test_parquet_cache_creation(self):
        """Verify that Parquet cache folder was created with the correct files."""
        cache_path = os.path.join(PARQUET_CACHE_DIR, "35")
        self.assertTrue(os.path.exists(cache_path), f"Parquet cache directory not found for Career 35: {cache_path}")
        
        # Check standard files exist
        dem_file = os.path.join(cache_path, "demografico.parquet")
        ind_file = os.path.join(cache_path, "indicadores.parquet")
        self.assertTrue(os.path.exists(dem_file), "demografico.parquet was not created")
        self.assertTrue(os.path.exists(ind_file), "indicadores.parquet was not created")

    def test_extraction_fields(self):
        """Verify that the extractor correctly extracts demographics and metadata."""
        data = extract_all_data(35)
        
        # Check metadata
        self.assertEqual(data["metadata"]["poblacion"], 28)
        self.assertEqual(data["metadata"]["muestra"], 14)
        self.assertEqual(data["metadata"]["cohorte"], "2025")
        
        # Check demographics
        dem = data["demographics"]
        self.assertEqual(dem["hombre_freq"], 13)
        self.assertEqual(dem["mujer_freq"], 1)
        self.assertAlmostEqual(dem["hombre_pct"], 0.9286, places=2)
        self.assertAlmostEqual(dem["mujer_pct"], 0.0714, places=2)
        
        # Check open question responses loaded
        self.assertIn("b8_nombres_emprendimientos", data["open_questions"])
        self.assertIn("f6_tematicas_de_capacitacion", data["open_questions"])
        self.assertIn("e2_que_agregarias_al_plan_curricular", data["open_questions"])
        
        # Verify that questions with naturally empty answers have 0 and those with answers are non-zero
        self.assertEqual(len(data["open_questions"]["b8_nombres_emprendimientos"]), 0)
        self.assertEqual(len(data["open_questions"]["f6_tematicas_de_capacitacion"]), 13)
        self.assertEqual(len(data["open_questions"]["e2_que_agregarias_al_plan_curricular"]), 1)

    def test_report_dry_run_update(self):
        """Runs the document updater on a test output file to verify modification correctness."""
        data = extract_all_data(35)
        
        # Temporary output docx in tests directory
        test_output_docx = os.path.join(os.path.dirname(__file__), "test_output_35.docx")
        if os.path.exists(test_output_docx):
            os.remove(test_output_docx)
            
        try:
            # Perform update on the test output file
            update_report(35, data, target_docx_path=test_output_docx)
            
            self.assertTrue(os.path.exists(test_output_docx), "Test output document was not saved")
            
            # Load and verify modified document tables
            doc = docx.Document(test_output_docx)
            
            # Verify Table 0 metadata
            t0 = doc.tables[0]
            cells_t0_poblacion = [row.cells[-1].text.strip() for row in t0.rows if "población" in "".join([c.text.lower() for c in row.cells])]
            cells_t0_muestra = [row.cells[-1].text.strip() for row in t0.rows if "encuestados" in "".join([c.text.lower() for c in row.cells])]
            
            self.assertEqual(cells_t0_poblacion[0], "28")
            self.assertEqual(cells_t0_muestra[0], "14")
            
            # Verify Table 4 demographics Male cell
            t4 = doc.tables[4]
            hombre_row = None
            for row in t4.rows:
                if row.cells[1].text.strip() == "Hombre":
                    hombre_row = row
                    break
            self.assertIsNotNone(hombre_row, "Male row not found in demographics table")
            self.assertEqual(hombre_row.cells[2].text.strip(), "13")
            self.assertEqual(hombre_row.cells[3].text.strip(), "92,86%")
            
        finally:
            # Cleanup temporary file
            if os.path.exists(test_output_docx):
                os.remove(test_output_docx)

if __name__ == "__main__":
    unittest.main()
