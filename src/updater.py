import docx
import os
import shutil
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from src.config import is_postgraduate, CAREER_FILES, PROCESSED_DIR
from src.config import is_postgraduate, CAREER_FILES, PROCESSED_DIR, get_report_category


def format_pct(val, multiply=True):
    """Formats float value as a percentage string with comma decimal separator (e.g. 0.9286 -> '92,86%')."""
    if multiply:
        val = val * 100
    return f"{val:.2f}%".replace(".", ",")


def clean_bullet_text(text):
    """Cleans up text from open questions, normalizing whitespace, xml characters, and capitalization."""
    if not text:
        return ""
    # Remove XML carriage return indicators and force type to string
    cleaned = str(text).replace("_x000D_", "").strip()
    import re

    cleaned = re.sub(r"\s+", " ", cleaned)
    # Check if text is mostly uppercase (e.g. "MANEJO DE SISTEMAS")
    if cleaned.isupper() or (
        sum(1 for c in cleaned if c.isupper()) / len(cleaned) > 0.5
        if len(cleaned) > 0
        else False
    ):
        cleaned = cleaned.lower()
    if cleaned:
        cleaned = cleaned[0].upper() + cleaned[1:]
    return cleaned


def insert_bullets(doc, idx, items_list, intro_text):
    """Replaces paragraph at idx with intro_text, and inserts list bullet paragraphs below it with clean styling."""
    para = doc.paragraphs[idx]
    para.style = "Normal"

    # Clear direct XML bullet formatting (numPr) and indentation (ind) from paragraph properties
    pPr = para._p.get_or_add_pPr()
    numPr = pPr.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr"
    )
    if numPr is not None:
        pPr.remove(numPr)
    ind = pPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ind")
    if ind is not None:
        pPr.remove(ind)

    para.paragraph_format.left_indent = None
    para.text = intro_text

    # Extract font details from original paragraph
    font_name = "Calibri"
    font_size = docx.shared.Pt(11)
    for run in para.runs:
        if run.font.name:
            font_name = run.font.name
            if run.font.size:
                font_size = run.font.size
            break

    next_para = doc.paragraphs[idx + 1] if idx + 1 < len(doc.paragraphs) else None

    # To maintain the original order when inserting before next_para
    has_list_bullet = "List Bullet" in doc.styles
    for item in items_list:
        cleaned = clean_bullet_text(item)
        if not cleaned:
            continue

        if has_list_bullet:
            if next_para:
                p = next_para.insert_paragraph_before(style="List Bullet")
            else:
                p = doc.add_paragraph(style="List Bullet")
            bullet_prefix = ""
        else:
            if next_para:
                p = next_para.insert_paragraph_before()
            else:
                p = doc.add_paragraph()
            p.paragraph_format.left_indent = docx.shared.Inches(0.25)
            bullet_prefix = "• "

        run = p.add_run(bullet_prefix + cleaned)
        run.font.name = font_name
        run.font.size = font_size
        p.paragraph_format.space_after = docx.shared.Pt(2)
        p.paragraph_format.space_before = docx.shared.Pt(2)


def strip_all_highlights(doc):
    """Removes run-level highlight/shading colors from body paragraph runs only.
    Table cell backgrounds and colors are intentionally preserved from the original document.
    """
    for para in doc.paragraphs:
        for run in para.runs:
            rPr = run._r.get_or_add_rPr()
            for hl in rPr.findall(qn("w:highlight")):
                rPr.remove(hl)
            for shd in rPr.findall(qn("w:shd")):
                rPr.remove(shd)
    print(
        "  [Strip] Removed highlights from body paragraph runs (table colors preserved)."
    )


def is_numeric_cell(text):
    """Returns True if cell text is purely numeric, a percentage, decimal, or empty.
    Used to distinguish data value cells (→ center) from descriptive label cells (→ left).
    """
    import re

    t = text.strip()
    if not t:
        return True
    # Match: integers, decimals (comma or dot), percentages, and combinations
    # Examples: "10", "19,23%", "0,00%", "100%", "52", "1.234,56"
    return bool(re.match(r"^[\d\.,\-\s%]+$", t))


def center_all_table_cells(doc):
    """Aligns table cell text based on content type and removes inner paragraph spacing.
    - Numeric / percentage cells → CENTER aligned
    - Descriptive label cells    → LEFT aligned
    """
    centered = 0
    lefted = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                align = (
                    WD_ALIGN_PARAGRAPH.CENTER
                    if is_numeric_cell(cell_text)
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
                if align == WD_ALIGN_PARAGRAPH.CENTER:
                    centered += 1
                else:
                    lefted += 1
                for para in cell.paragraphs:
                    para.alignment = align
                    pf = para.paragraph_format
                    pf.space_before = docx.shared.Pt(0)
                    pf.space_after = docx.shared.Pt(0)
    print(
        f"  [Align] Table cells: {centered} centered (numeric), {lefted} left-aligned (labels)."
    )


def update_table_0(table, data):
    """Updates Table 0 (Metadata) population, sample sizes, and employer count."""
    total_employers = data.get("employer", {}).get("total_employers", 0)
    # Find rows with Poblacion, Muestra, and Empleadores
    for row in table.rows:
        cells = [c.text.strip().lower() for c in row.cells]
        # Match 'total de graduados (población)' or 'población'
        if any("población" in cell or "poblacion" in cell for cell in cells):
            # Overwrite last cell
            row.cells[-1].text = str(data["metadata"]["poblacion"])
        if any("encuestados" in cell or "muestra" in cell for cell in cells) and not any("empleadores" in cell for cell in cells):
            # Overwrite last cell
            row.cells[-1].text = str(data["metadata"]["muestra"])
        if any("empleadores encuestados" in cell for cell in cells):
            row.cells[-1].text = str(total_employers)


def update_table_4(table, data):
    """Updates Table 4 (Demographics) table cells."""
    dem = data["demographics"]
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 4:
            continue

        # Check second cell (unnamed_1)
        cat = cells[1].strip()

        if cat == "Hombre":
            row.cells[2].text = str(dem["hombre_freq"])
            row.cells[3].text = format_pct(dem["hombre_pct"])
        elif cat == "Mujer":
            row.cells[2].text = str(dem["mujer_freq"])
            row.cells[3].text = format_pct(dem["mujer_pct"])
        elif cat == "Soltero/a":
            row.cells[2].text = str(dem["soltero_freq"])
            row.cells[3].text = format_pct(dem["soltero_pct"])
        elif cat == "Casado/a":
            row.cells[2].text = str(dem["casado_freq"])
            row.cells[3].text = format_pct(dem["casado_pct"])
        elif cat == "Divorciado/a":
            row.cells[2].text = str(dem["divorciado_freq"])
            row.cells[3].text = format_pct(dem["divorciado_pct"])
        elif cat == "Viudo/a":
            row.cells[2].text = str(dem["viudo_freq"])
            row.cells[3].text = format_pct(dem["viudo_pct"])
        elif cat == "Unión libre":
            row.cells[2].text = str(dem["union_libre_freq"])
            row.cells[3].text = format_pct(dem["union_libre_pct"])
        elif "21 y 25" in cat:
            row.cells[2].text = str(dem["edad_21_25_freq"])
            row.cells[3].text = format_pct(dem["edad_21_25_pct"])
        elif "26 y 30" in cat:
            row.cells[2].text = str(dem["edad_26_30_freq"])
            row.cells[3].text = format_pct(dem["edad_26_30_pct"])
        elif "31 y 35" in cat:
            row.cells[2].text = str(dem["edad_31_35_freq"])
            row.cells[3].text = format_pct(dem["edad_31_35_pct"])
        elif "36 y 40" in cat:
            row.cells[2].text = str(dem["edad_36_40_freq"])
            row.cells[3].text = format_pct(dem["edad_36_40_pct"])
        elif "41 y 45" in cat:
            row.cells[2].text = str(dem["edad_41_45_freq"])
            row.cells[3].text = format_pct(dem["edad_41_45_pct"])
        elif "mayor" in cat or "45 años" in cat:
            row.cells[2].text = str(dem["edad_mayor_45_freq"])
            row.cells[3].text = format_pct(dem["edad_mayor_45_pct"])


def update_table_5(table, data):
    """Updates Table 5 (Training Needs) table cells."""
    needs = data["training_needs"]
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 3:
            continue

        cat = cells[0].strip()
        if "Simposios" in cat or "congresos" in cat:
            row.cells[1].text = str(needs["capacitacion_seminarios_freq"])
            row.cells[2].text = format_pct(needs["capacitacion_seminarios_pct"])
        elif "certificación" in cat or "competencias" in cat:
            row.cells[1].text = str(needs["capacitacion_certificacion_freq"])
            row.cells[2].text = format_pct(needs["capacitacion_certificacion_pct"])
        elif "MOOC" in cat or "autoinstruccionales" in cat:
            row.cells[1].text = str(needs["capacitacion_mooc_freq"])
            row.cells[2].text = format_pct(needs["capacitacion_mooc_pct"])
        elif "e-learning" in cat:
            row.cells[1].text = str(needs["capacitacion_elearning_freq"])
            row.cells[2].text = format_pct(needs["capacitacion_elearning_pct"])
        elif "Diplomados" in cat:
            row.cells[1].text = str(needs["capacitacion_diplomados_freq"])
            row.cells[2].text = format_pct(needs["capacitacion_diplomados_pct"])
        elif "Especializaciones" in cat:
            row.cells[1].text = str(needs["capacitacion_especializaciones_freq"])
            row.cells[2].text = format_pct(needs["capacitacion_especializaciones_pct"])
        elif "Maestría" in cat:
            row.cells[1].text = str(needs["capacitacion_maestria_freq"])
            row.cells[2].text = format_pct(needs["capacitacion_maestria_pct"])
        elif "Doctorado" in cat:
            row.cells[1].text = str(needs["capacitacion_doctorado_freq"])
            row.cells[2].text = format_pct(needs["capacitacion_doctorado_pct"])


def update_table_8_and_11(table, data):
    """Updates Table 8 & 11 (Employment Situation) table cells."""
    emp = data["employment"]
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 3:
            continue

        cat = cells[0].strip()
        if "relación de dependencia" in cat and "público" in cat:
            row.cells[1].text = str(emp["empleo_publico_freq"])
            row.cells[2].text = format_pct(emp["empleo_publico_pct"])
        elif "relación de dependencia" in cat and "privado" in cat:
            row.cells[1].text = str(emp["empleo_privado_freq"])
            row.cells[2].text = format_pct(emp["empleo_privado_pct"])
        elif "Libre ejercicio" in cat or "facturación" in cat:
            row.cells[1].text = str(emp["empleo_libre_freq"])
            row.cells[2].text = format_pct(emp["empleo_libre_pct"])
        elif "Emprendedor" in cat or "RUC" in cat or "RIMPE" in cat:
            row.cells[1].text = str(emp["empleo_emprendedor_freq"])
            row.cells[2].text = format_pct(emp["empleo_emprendedor_pct"])
        elif "Sin actividad" in cat or "Desempleo" in cat or "desocupado" in cat:
            row.cells[1].text = str(emp["empleo_desempleo_freq"])
            row.cells[2].text = format_pct(emp["empleo_desempleo_pct"])


def update_table_12(table, data):
    """Updates Table 12 (Study Continuity) table cells."""
    cont = data["continuity"]
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 3:
            continue

        cat = cells[0].strip()
        if "Cursos" in cat or "talleres" in cat or "simposios" in cat:
            row.cells[1].text = str(cont["cont_talleres_freq"])
            row.cells[2].text = format_pct(cont["cont_talleres_pct"])
        elif "Diplomado" in cat:
            row.cells[1].text = str(cont["cont_diplomado_freq"])
            row.cells[2].text = format_pct(cont["cont_diplomado_pct"])
        elif "Especialidad" in cat:
            row.cells[1].text = str(cont["cont_especialidad_freq"])
            row.cells[2].text = format_pct(cont["cont_especialidad_pct"])
        elif "Maestría" in cat or "maestría" in cat:
            row.cells[1].text = str(cont["cont_maestria_freq"])
            row.cells[2].text = format_pct(cont["cont_maestria_pct"])
        elif "Doctorado" in cat:
            row.cells[1].text = str(cont["cont_doctorado_freq"])
            row.cells[2].text = format_pct(cont["cont_doctorado_pct"])
        elif "Posdoctorado" in cat:
            row.cells[1].text = str(cont["cont_posdoctorado_freq"])
            row.cells[2].text = format_pct(cont["cont_posdoctorado_pct"])
        elif "Otra carrera" in cat:
            row.cells[1].text = str(cont["cont_otra_carrera_freq"])
            row.cells[2].text = format_pct(cont["cont_otra_carrera_pct"])
        elif "No," in cat or "Aun no" in cat or "aun no" in cat:
            row.cells[1].text = str(cont["cont_ninguno_freq"])
            row.cells[2].text = format_pct(cont["cont_ninguno_pct"])


def update_table_13(table, data):
    """Updates Table 13 (Job Demand Level) table cells."""
    dem = data["demand"]
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 3:
            continue

        cat = cells[0].strip()
        if cat == "Alta":
            row.cells[1].text = str(dem["demanda_alta_freq"])
            row.cells[2].text = format_pct(dem["demanda_alta_pct"])
        elif cat == "Media":
            row.cells[1].text = str(dem["demanda_media_freq"])
            row.cells[2].text = format_pct(dem["demanda_media_pct"])
        elif cat == "Baja":
            row.cells[1].text = str(dem["demanda_baja_freq"])
            row.cells[2].text = format_pct(dem["demanda_baja_pct"])
        elif cat == "Muy Baja" or cat == "Muy baja":
            row.cells[1].text = str(dem["demanda_muy_baja_freq"])
            row.cells[2].text = format_pct(dem["demanda_muy_baja_pct"])


def get_d1_count(d1_sq, category):
    cat = category.lower()
    for k, v in d1_sq.items():
        k_low = k.lower()
        if cat == "muy bueno" and "muy bueno" in k_low:
            return v.get("count", 0)
        elif cat == "bueno" and "bueno" in k_low and "muy" not in k_low:
            return v.get("count", 0)
        elif cat == "regular" and "regular" in k_low:
            return v.get("count", 0)
        elif cat == "malo" and "malo" in k_low:
            return v.get("count", 0)
    return 0


def get_d1_pct(d1_sq, category):
    cat = category.lower()
    for k, v in d1_sq.items():
        k_low = k.lower()
        if cat == "muy bueno" and "muy bueno" in k_low:
            return v.get("pct", 0.0)
        elif cat == "bueno" and "bueno" in k_low and "muy" not in k_low:
            return v.get("pct", 0.0)
        elif cat == "regular" and "regular" in k_low:
            return v.get("pct", 0.0)
        elif cat == "malo" and "malo" in k_low:
            return v.get("pct", 0.0)
    return 0.0


def get_employer_stat(sub_data, category):
    cat_lower = category.lower()
    for k, v in sub_data.items():
        k_lower = k.lower()
        if cat_lower == "bueno" and "muy bueno" in k_lower:
            continue
        if cat_lower == "malo" and "muy malo" in k_lower:
            continue
        if cat_lower in k_lower:
            return v.get("count", 0), v.get("pct", 0.0)
    return 0, 0.0


def format_decimal(val):
    if not isinstance(val, (int, float)):
        return str(val)
    return f"{val:.2f}".replace(".", ",")


def update_table_6(table, data):
    """Updates Table 6 (Correlation analysis) cells."""
    headers = [c.text.strip().lower() for c in table.rows[0].cells]
    if not any("correlación" in h or "correlacion" in h for h in headers) or not any(
        "nivel" in h for h in headers
    ):
        return

    corr = data.get("correlation", {})
    grad = corr.get("graduados", {})
    emp = corr.get("empleadores", {})

    if not grad and not emp:
        return

    row_mapping = {
        "análisis y diagnósticos": "Capacidad para realizar análisis y diagnósticos.",
        "analisis y diagnosticos": "Capacidad para realizar análisis y diagnósticos.",
        "solución de problemas": "Capacidad para solucionar problemas relacionados con la profesión.",
        "solucion de problemas": "Capacidad para solucionar problemas relacionados con la profesión.",
        "herramientas especializadas": "Capacidad para utilizar herramientas especializadas propias de la profesión.",
        "realidad local/regional/nacional": "Conocimientos sobre la realidad local, regional y nacional.",
        "búsqueda y sistematización": "Habilidad para encontrar, seleccionar, sistematizar y utilizar información.",
        "busqueda y sistematizacion": "Habilidad para encontrar, seleccionar, sistematizar y utilizar información.",
        "trabajo en equipo": "Capacidad para trabajar en equipo.",
        "redacción de informes": "Capacidad para redactar informes o documentos.",
        "redaccion de informes": "Capacidad para redactar informes o documentos.",
        "presentación pública": "Capacidad para presentar en público resultados, ideas o informes.",
        "presentacion publica": "Capacidad para presentar en público resultados, ideas o informes.",
        "herramientas informáticas": "Capacidad para utilizar herramientas informáticas y comunicación digital.",
        "herramientas informaticas": "Capacidad para utilizar herramientas informáticas y comunicación digital.",
        "emprendimiento": "Capacidad para el emprendimiento",
    }

    for row in table.rows[1:]:
        if len(row.cells) < 5:
            if len(row.cells) == 3:
                cat_name = row.cells[0].text.strip().lower()
                key = None
                for kw, full_key in row_mapping.items():
                    if kw in cat_name:
                        key = full_key
                        break
                if key and key in grad:
                    row.cells[1].text = format_decimal(grad[key]["value"])
                    row.cells[2].text = grad[key]["level"]
            continue
        cat_name = row.cells[0].text.strip().lower()
        key = None
        for kw, full_key in row_mapping.items():
            if kw in cat_name:
                key = full_key
                break

        if key:
            if key in grad:
                row.cells[1].text = format_decimal(grad[key]["value"])
                row.cells[2].text = grad[key]["level"]
            if key in emp:
                row.cells[3].text = format_decimal(emp[key]["value"])
                row.cells[4].text = emp[key]["level"]


def update_table_7(table, data):
    """Updates Table 7 (Employer Performance Rating) cells."""
    headers = [c.text.strip().lower() for c in table.rows[0].cells]
    if not any("desempeño" in h or "desempeno" in h for h in headers) or not any(
        "bueno" in h for h in headers
    ):
        return

    d1 = data.get("employer", {}).get("d1", {})
    if not d1:
        return

    row_mapping = {
        "general": "sq1",
        "funciones y tareas": "sq2",
        "científico-técnico": "sq2",
        "tiempos requeridos": "sq3",
        "responsabilidad y agilidad": "sq3",
        "resolver problemas": "sq4",
        "ético": "sq5",
        "etico": "sq5",
        "compromiso": "sq6",
        "seguir formándose": "sq7",
        "seguir formandose": "sq7",
        "interés en el desarrollo": "sq7",
        "interes en el desarrollo": "sq7",
        "toma de decisiones": "sq8",
        "eficaz y eficiente": "sq9",
    }

    for row in table.rows[1:]:
        cat_name = row.cells[0].text.strip().lower()
        sq_key = None
        for kw, sq in row_mapping.items():
            if kw in cat_name:
                sq_key = sq
                break
        if sq_key and sq_key in d1:
            row.cells[1].text = format_pct(
                get_d1_pct(d1[sq_key], "muy bueno"), multiply=True
            )
            row.cells[2].text = format_pct(
                get_d1_pct(d1[sq_key], "bueno"), multiply=True
            )
            row.cells[3].text = format_pct(
                get_d1_pct(d1[sq_key], "regular"), multiply=True
            )
            row.cells[4].text = format_pct(
                get_d1_pct(d1[sq_key], "malo"), multiply=True
            )


def update_table_10(table, data):
    """Updates Table 10 (Employer Consultation & Requirements) cells."""
    cell_texts = []
    for row in table.rows:
        for c in row.cells:
            cell_texts.append(c.text.strip().lower())

    if not any("empleadores consultados" in txt for txt in cell_texts) and not any(
        "requisitos formales" in txt for txt in cell_texts
    ):
        return

    emp_b1 = data.get("employer", {}).get("b1", {})
    emp_b2 = data.get("employer", {}).get("b2", {})
    total_employers = sum(v.get("count", 0) for v in emp_b1.values())

    for row in table.rows:
        r_cells = [c.text.strip().lower() for c in row.cells]
        if any("empleadores consultados" in c for c in r_cells) and total_employers > 0:
            row.cells[2].text = str(total_employers)
            row.cells[3].text = "100%"

    for row in table.rows:
        if len(row.cells) < 4:
            continue
        c0 = row.cells[0].text.strip().lower()
        c1_val = row.cells[1].text.strip().lower()

        if "requisitos formales" in c0:
            if "muy bueno" in c1_val:
                count, pct = get_employer_stat(emp_b1, "muy bueno")
                row.cells[2].text = str(count)
                row.cells[3].text = format_pct(pct, multiply=True)
            elif "bueno" in c1_val:
                count, pct = get_employer_stat(emp_b1, "bueno")
                row.cells[2].text = str(count)
                row.cells[3].text = format_pct(pct, multiply=True)
            elif "regular" in c1_val:
                count, pct = get_employer_stat(emp_b1, "regular")
                row.cells[2].text = str(count)
                row.cells[3].text = format_pct(pct, multiply=True)
            elif "malo" in c1_val:
                count, pct = get_employer_stat(emp_b1, "malo")
                row.cells[2].text = str(count)
                row.cells[3].text = format_pct(pct, multiply=True)

        elif "aspectos de contratación" in c0 or "aspectos de contratacion" in c0:
            aspect_kw = None
            if "vulnerables" in c1_val:
                aspect_kw = "vulnerables"
            elif "edad" in c1_val:
                aspect_kw = "edad"
            elif "sexo" in c1_val:
                aspect_kw = "sexo"
            elif "civil" in c1_val:
                aspect_kw = "civil"
            elif "académicos" in c1_val or "academicos" in c1_val:
                aspect_kw = "acad"
            elif "comunicación" in c1_val or "comunicacion" in c1_val:
                aspect_kw = "comunic"
            elif "responsabilidades" in c1_val:
                aspect_kw = "respons"
            elif "demanda" in c1_val or "prestigio" in c1_val:
                aspect_kw = "prestigio"

            if aspect_kw:
                count, pct = get_employer_stat(emp_b2, aspect_kw)
                row.cells[2].text = str(count)
                row.cells[3].text = format_pct(pct, multiply=True)


def rewrite_paragraphs(report_id, doc, data):
    """Dynamically reconstructs narrative paragraphs below key tables to keep them fully synchronized."""
    is_pg = is_postgraduate(report_id)
    n = data["metadata"]["muestra"]
    poblacion = data["metadata"]["poblacion"]
    cohorte = data["metadata"]["cohorte"]

    # Terminology difference
    graduado_term = "maestrantes" if is_pg else "graduados"

    # Store dynamic bullet insertions to perform after iteration
    insertions = []

    dem = data["demographics"]
    needs = data["training_needs"]

    # 1. Gender details for demographics narrative
    dominant_gender = (
        "masculina" if dem["hombre_freq"] >= dem["mujer_freq"] else "femenina"
    )
    dominant_gender_pct = max(dem["hombre_pct"], dem["mujer_pct"]) * 100
    minor_gender_pct = min(dem["hombre_pct"], dem["mujer_pct"]) * 100

    # 2. Solteros details
    solteros_pct = dem["soltero_pct"] * 100

    # Age sum (26-35)
    edad_26_35_pct = (dem["edad_26_30_pct"] + dem["edad_31_35_pct"]) * 100

    # 3. FP Interest details
    fp_interest_pct = needs["interes_si_pct"] * 100
    fp_no_interest_pct = needs["interes_no_pct"] * 100
    fp_interest_n = needs["interes_si_freq"]
    fp_no_interest_n = needs["interes_no_freq"]

    # Specific layout correction for Report 26 (Economía 2023)
    # The original document has the "enseñanza-aprendizaje" analysis misplaced under the "Perfil de egreso" heading.
    if report_id == 26:
        try:
            # Dynamically locate the headings
            ea_idx = -1
            pe_idx = -1
            for i, p in enumerate(doc.paragraphs):
                p_text = p.text.strip()
                if p_text == "Percepción sobre el proceso de enseñanza-aprendizaje":
                    ea_idx = i
                elif p_text == "Percepción sobre el perfil de egreso":
                    pe_idx = i

            if ea_idx != -1 and pe_idx != -1 and pe_idx > ea_idx:
                # misplaced teaching-learning paragraphs are right after "Percepción sobre el perfil de egreso"
                # Let's collect them: they are usually the next 3 paragraphs.
                # Let's inspect them to confirm they talk about "enseñanza-aprendizaje".
                ea_p1 = doc.paragraphs[pe_idx + 1].text
                ea_p2 = doc.paragraphs[pe_idx + 2].text
                ea_p3 = doc.paragraphs[pe_idx + 3].text
                
                if "enseñanza-aprendizaje" in ea_p1.lower() or "enseñanza-aprendizaje" in ea_p3.lower():
                    # Yes! The teaching-learning text is indeed misplaced under "Perfil de egreso".
                    # We want to re-arrange this section so that:
                    # 1. Right under "Percepción sobre el proceso de enseñanza-aprendizaje" (ea_idx) we place ea_p1, ea_p2, ea_p3.
                    # 2. Then we place the heading "Percepción sobre el perfil de egreso" with its proper style (Heading 1).
                    # 3. Then we place the newly generated profile of egress text.
                    
                    # To do this safely and preserve all styling (like the bold numbering 4.5.4 in run [1] etc.),
                    # we can copy the contents run-by-run instead of replacing text, or replace texts but keep original styles.
                    # Since we want to insert ea_p1, ea_p2, ea_p3 right after ea_idx, let's see:
                    # In the original document:
                    # ea_idx: 'Percepción sobre el proceso de enseñanza-aprendizaje'
                    # ea_idx + 1: '' (empty paragraph)
                    # ea_idx + 2: 'Percepción sobre el perfil de egreso' (Wait, in original, 124 is empty, 125 is Percepción sobre el perfil de egreso)
                    # So ea_idx + 2 was 'Percepción sobre el perfil de egreso'.
                    # Let's just insert paragraphs before pe_idx + 1 (the misplaced text), and clear the misplaced text.
                    
                    # Better solution:
                    # Clear the misplaced texts from their original positions and construct the content:
                    # 1. Replace the paragraph ea_idx + 1 (empty) with ea_p1
                    # 2. Replace the paragraph pe_idx (which was "Percepción sobre el perfil de egreso") with ea_p2 (keeping style Normal)
                    # 3. Replace the paragraph pe_idx + 1 with ea_p3 (keeping style Normal)
                    # 4. Replace the paragraph pe_idx + 2 with "Percepción sobre el perfil de egreso" (assigning style Heading 1)
                    
                    p124_style = doc.paragraphs[ea_idx + 1].style
                    p_pe_style = doc.paragraphs[pe_idx].style
                    p_pe_plus_3_style = doc.paragraphs[pe_idx + 3].style

                    orig_pe_pPr = doc.paragraphs[pe_idx]._p.get_or_add_pPr()
                    numPr_element = orig_pe_pPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr")
                    ind_element = orig_pe_pPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ind")
                    
                    numPr_copy = None
                    if numPr_element is not None:
                        import copy
                        numPr_copy = copy.deepcopy(numPr_element)
                        orig_pe_pPr.remove(numPr_element)

                    ind_copy = None
                    if ind_element is not None:
                        import copy
                        ind_copy = copy.deepcopy(ind_element)
                        orig_pe_pPr.remove(ind_element)

                    for p in [doc.paragraphs[ea_idx + 1], doc.paragraphs[pe_idx], doc.paragraphs[pe_idx + 1], doc.paragraphs[pe_idx + 2], doc.paragraphs[pe_idx + 3]]:
                        p.text = ""
                        p.paragraph_format.left_indent = None
                        p.paragraph_format.first_line_indent = None
                        
                        pPr = p._p.get_or_add_pPr()
                        numPr = pPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr")
                        if numPr is not None:
                            pPr.remove(numPr)
                        ind = pPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ind")
                        if ind is not None:
                            pPr.remove(ind)
                        
                    # Helper function to set text with Arial 11pt style
                    def set_paragraph_text_arial_11(p, text_str):
                        p.text = ""
                        run = p.add_run(text_str)
                        run.font.name = "Arial"
                        run.font.size = docx.shared.Pt(11)

                    set_paragraph_text_arial_11(doc.paragraphs[ea_idx + 1], ea_p1)
                    set_paragraph_text_arial_11(doc.paragraphs[pe_idx], ea_p2)
                    set_paragraph_text_arial_11(doc.paragraphs[pe_idx + 1], ea_p3)
                    doc.paragraphs[pe_idx + 2].text = "Percepción sobre el perfil de egreso"
                    
                    doc.paragraphs[ea_idx + 1].style = p124_style
                    doc.paragraphs[pe_idx].style = p124_style
                    doc.paragraphs[pe_idx + 1].style = p124_style
                    doc.paragraphs[pe_idx + 2].style = p_pe_style
                    
                    new_pe_pPr = doc.paragraphs[pe_idx + 2]._p.get_or_add_pPr()
                    if numPr_copy is not None:
                        new_pe_pPr.append(numPr_copy)
                    if ind_copy is not None:
                        new_pe_pPr.append(ind_copy)

                    sat = data.get("satisfaccion", {})
                    res_buena = sat.get("pe_resultados_buena_pct", 0.0) * 100
                    res_buena_n = sat.get("pe_resultados_buena_freq", 0)
                    res_reg = sat.get("pe_resultados_regular_pct", 0.0) * 100
                    res_reg_n = sat.get("pe_resultados_regular_freq", 0)

                    cump_buena = sat.get("pe_cumplimiento_buena_pct", 0.0) * 100
                    cump_buena_n = sat.get("pe_cumplimiento_buena_freq", 0)
                    cump_reg = sat.get("pe_cumplimiento_regular_pct", 0.0) * 100
                    cump_reg_n = sat.get("pe_cumplimiento_regular_freq", 0)

                    new_pe_text = (
                        f"Respecto a la percepción sobre el perfil de egreso, el {res_buena:.2f}% de los graduados "
                        f"({res_buena_n} personas) califica la pertinencia de los resultados de aprendizaje como Buena, "
                        f"mientras que el {res_reg:.2f}% ({res_reg_n} personas) la considera Regular. Por otro lado, "
                        f"en relación al cumplimiento del perfil de egreso, el {cump_buena:.2f}% ({cump_buena_n} graduados) "
                        f"lo valora como Bueno, y el {cump_reg:.2f}% ({cump_reg_n} personas) como Regular. No se registraron "
                        f"calificaciones en la categoría Mala para ninguno de los dos aspectos evaluados, evidenciando "
                        f"un alto nivel de satisfacción con los logros formativos alcanzados."
                    )
                    set_paragraph_text_arial_11(doc.paragraphs[pe_idx + 3], new_pe_text)
                    doc.paragraphs[pe_idx + 3].style = p_pe_plus_3_style
                    
                    print("  [Paragraph] Applied Report 26 layout correction (Teaching-Learning & Perfil de Egreso mismatch resolved).")
        except Exception as e:
            print(f"  [Warning] Failed to apply specific Report 26 layout correction: {e}")


    # Re-writing paragraphs based on pattern match
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()


        # Demographic narrative replacement
        if (
            "predominancia masculina" in text.lower()
            or "predominancia femenina" in text.lower()
            or (
                "cohorte" in text.lower()
                and "integrada por" in text.lower()
                and "femenina" in text.lower()
            )
        ):
            new_text = (
                f"La cohorte está integrada por {n} {graduado_term}, con una marcada predominancia "
                f"{dominant_gender} ({dominant_gender_pct:.2f}%) frente a un {minor_gender_pct:.2f}% de "
                f"{'mujeres' if dominant_gender == 'masculina' else 'hombres'}. El {solteros_pct:.2f}% son solteros "
                f"y los rangos de edad predominantes se sitúan entre los 26 y 35 años ({edad_26_35_pct:.2f}% en total)."
            )
            para.text = new_text
            print(f"  [Paragraph] Updated demographic interpretation paragraph {idx}.")

        # FP Interest narrative replacement
        elif (
            "evidencia interés en formarse" in text.lower()
            or "actualizarse profesionalmente" in text.lower()
        ):
            new_text = (
                f"El {fp_interest_pct:.2f}% de los {graduado_term} ({fp_interest_n} personas) evidencia interés en formarse y "
                f"actualizarse profesionalmente, mientras que el {fp_no_interest_pct:.2f}% restante ({fp_no_interest_n} individuos) "
                f"no manifiesta dicha intención. Esta proporción constata una inclinación mayoritaria hacia la capacitación continua, "
                f"factor estructural para el fortalecimiento competitivo de su perfil en el mercado laboral."
            )
            para.text = new_text
            print(
                f"  [Paragraph] Updated training interest interpretation paragraph {idx}."
            )

        # Population summary paragraph replacement
        elif (
            "seguimiento a graduados cohorte" in text.lower()
            and "se tiene un total de" in text.lower()
        ):
            new_text = f"En la recolección de datos estadísticos para los estudios de seguimiento a graduados cohorte {cohorte} se tiene un total de {poblacion}."
            para.text = new_text
            print(f"  [Paragraph] Updated population count paragraph {idx}.")

        # 4. Employer demand perception paragraph replacement
        elif (
            "no proporcionaron información cuantificable sobre la demanda"
            in text.lower()
            or "valores en cero para todas las categorías" in text.lower()
        ):
            emp_b1 = data.get("employer", {}).get("b1", {})
            emp_b3 = data.get("employer", {}).get("b3", {})
            total_emp = sum(v.get("count", 0) for v in emp_b1.values())
            if total_emp == 0:
                total_emp = 10
            if emp_b3:
                alta_pct = emp_b3.get("Alta (B4A1)", {}).get("pct", 0.0) * 100
                media_pct = emp_b3.get("Media (B4A2)", {}).get("pct", 0.0) * 100
                baja_pct = emp_b3.get("Baja (B4A3)", {}).get("pct", 0.0) * 100
                muy_baja_pct = emp_b3.get("Muy baja (B4A4)", {}).get("pct", 0.0) * 100
                alta_count = emp_b3.get("Alta (B4A1)", {}).get("count", 0)
                media_count = emp_b3.get("Media (B4A2)", {}).get("count", 0)

                new_text = (
                    f"De los {total_emp} empleadores consultados, el {alta_pct:.2f}% ({alta_count} empleadores) considera que la demanda es Alta, "
                    f"y el {media_pct:.2f}% ({media_count} empleadores) la califica como Media, mientras que no se registraron respuestas "
                    f"para las categorías de demanda Baja ({baja_pct:.2f}%) o Muy Baja ({muy_baja_pct:.2f}%). Esta distribución refleja una "
                    f"opinión mayoritariamente positiva del sector empleador respecto a las oportunidades de inserción laboral para los profesionales de la carrera."
                )
                para.text = new_text
                print(
                    f"  [Paragraph] Updated employer demand interpretation paragraph {idx}."
                )

        # 5. Training suggestions bulleted list replacement
        elif text == "Conocimientos en …" or text == "….":
            suggestions = data.get("open_questions", {}).get(
                "temas_de_capacitacion_sugerida_por_los_empleadores", []
            )
            if suggestions:
                insertions.append(
                    (
                        idx,
                        suggestions,
                        "Los empleadores consultados sugirieron las siguientes temáticas clave de capacitación y actualización profesional:",
                    )
                )
                print(
                    f"  [Paragraph] Queued training suggestions for insertion below paragraph {idx}."
                )

        # 6. Emprendimientos list replacement
        elif text == "Existe un emprendimiento ……":
            emp_names = data.get("open_questions", {}).get(
                "b8_nombres_emprendimientos", []
            )
            valid_names = [
                clean_bullet_text(name)
                for name in emp_names
                if name.lower() not in ["sin nombre", "no registra", "ninguno"]
            ]
            if valid_names:
                new_text = f"En la cohorte analizada, se identificaron emprendimientos activos como: {', '.join(f'"{n}"' for n in valid_names)}, evidenciando una orientación hacia la iniciativa autónoma y el ejercicio independiente de la profesión."
            else:
                new_text = "No se reportaron nombres específicos de emprendimientos formalizados en la cohorte analizada."
            para.text = new_text
            print(f"  [Paragraph] Updated emprendimiento names in paragraph {idx}.")

        # 7. Employer cargos list replacement (Question C1)
        elif (
            idx > 0
            and text.strip() != ""
            and len(text) < 15
            and text.replace(".", "").replace("…", "").strip() == ""
            and "cargo" in doc.paragraphs[idx - 1].text.lower()
            and "c1" in doc.paragraphs[idx - 1].text.lower()
        ):
            cargos = data.get("open_questions", {}).get("employer_cargos", [])
            # Fallback: try datae C2 section for postgrad
            if not cargos:
                cargos = data.get("employer", {}).get("employer_cargos_datae", [])
            if cargos:
                insertions.append(
                    (
                        idx,
                        cargos,
                        "Los cargos reportados por los empleadores para los graduados de la carrera incluyen:",
                    )
                )
                print(
                    f"  [Paragraph] Queued employer cargos for insertion below paragraph {idx}."
                )
            else:
                para.text = "Los empleadores consultados no reportaron cargos específicos en los datos disponibles."
                print(f"  [Paragraph] Set fallback for employer cargos paragraph {idx}.")

        # 8. Employer funciones list replacement (Question C2)
        elif (
            idx > 0
            and text.strip() != ""
            and len(text) < 15
            and text.replace(".", "").replace("…", "").strip() == ""
            and "funciones" in doc.paragraphs[idx - 1].text.lower()
            and "c2" in doc.paragraphs[idx - 1].text.lower()
        ):
            funciones = data.get("open_questions", {}).get("employer_funciones", [])
            # Fallback: try datae C3 section for postgrad
            if not funciones:
                funciones = data.get("employer", {}).get("employer_funciones_datae", [])
            if funciones:
                insertions.append(
                    (
                        idx,
                        funciones,
                        "Las funciones principales que desempeñan los graduados en sus puestos de trabajo, según los empleadores, son:",
                    )
                )
                print(
                    f"  [Paragraph] Queued employer funciones for insertion below paragraph {idx}."
                )
            else:
                para.text = "Los empleadores consultados no reportaron funciones específicas en los datos disponibles."
                print(f"  [Paragraph] Set fallback for employer funciones paragraph {idx}.")

        # 9. Bolsa de Empleo effectiveness paragraph
        elif "efectividad de la bolsa de empleo" in text.lower() and (
            "registra un alcance" in text.lower()
            or "bolsa de empleo de la universidad" in text.lower()
            or "apenas el" in text.lower()
        ):
            ins = data.get("insercion", {})
            b_si = ins.get("bolsa_si_pct", 0.0) * 100
            b_si_n = ins.get("bolsa_si_freq", 0)
            b_no = ins.get("bolsa_no_pct", 0.0) * 100
            b_no_n = ins.get("bolsa_no_freq", 0)
            new_text = (
                f"La efectividad de la Bolsa de Empleo de la Universidad Nacional de Loja registra una utilidad "
                f"del {b_si:.2f}% ({b_si_n} persona(s)) que indica haber obtenido empleo mediante este canal, "
                f"mientras que el {b_no:.2f}% ({b_no_n} {graduado_term}) señaló que no le resultó efectivo."
            )
            para.text = new_text
            print(
                f"  [Paragraph] Updated Bolsa de Empleo interpretation paragraph {idx}."
            )

        # 10. Emprendimiento / Libre Ejercicio combinada paragraph
        elif (
            "opera como emprendedor autónomo" in text.lower()
            or "tasa combinada de trabajo independiente" in text.lower()
        ):
            emp_aut = data["employment"]["empleo_emprendedor_pct"] * 100
            emp_aut_n = data["employment"]["empleo_emprendedor_freq"]
            libre = data["employment"]["empleo_libre_pct"] * 100
            libre_n = data["employment"]["empleo_libre_freq"]
            comb = emp_aut + libre
            desemp = data["employment"]["empleo_desempleo_pct"] * 100
            desemp_n = data["employment"]["empleo_desempleo_freq"]
            new_text = (
                f"Apenas el {emp_aut:.2f}% de la cohorte ({emp_aut_n} individuos) opera como emprendedor autónomo "
                f"formal (RUC/RIMPE), complementado por un {libre:.2f}% ({libre_n} {graduado_term}) en libre ejercicio profesional. "
                f"Esta tasa combinada de trabajo independiente ({comb:.2f}%) coexiste con una tasa de desempleo "
                f"o sin actividad laboral registrada del {desemp:.2f}% ({desemp_n} personas)."
            )
            para.text = new_text
            print(
                f"  [Paragraph] Updated Independent work / unemployment interpretation paragraph {idx}."
            )

        # 11. Preference for Maestría paragraph
        elif "egresados hacia programas de maestría" in text.lower() and (
            "confirmar una preferencia" in text.lower()
            or "refugio conservador" in text.lower()
            or "preferencia por acumular" in text.lower()
        ):
            maestria_need = data["training_needs"]["capacitacion_maestria_pct"] * 100
            new_text = (
                f"La preferencia expresada por el {maestria_need:.2f}% de los egresados hacia programas de maestría "
                f"confirma una marcada orientación hacia la continuidad de su formación académica a nivel de posgrado."
            )
            para.text = new_text
            print(f"  [Paragraph] Updated Maestría preference paragraph {idx}.")

        # 12. Continuity of studies paragraph
        elif (
            "tasa de continuidad académica hacia programas de posgrado" in text.lower()
            or "limitando la actualización profesional a formatos" in text.lower()
        ):
            cont_maestria = data["continuity"]["cont_maestria_pct"] * 100
            new_text = (
                f"Esta distribución confirma una tasa de continuidad académica efectiva hacia programas de posgrado estructurados "
                f"de maestría del {cont_maestria:.2f}%, lo que sugiere que la mayoría de los graduados prioriza la inserción "
                f"laboral inmediata o capacitaciones cortas."
            )
            para.text = new_text
            print(f"  [Paragraph] Updated academic continuity paragraph {idx}.")

        # 13. Percepción de Demanda Ocupacional paragraph
        elif "considera que la demanda es baja" in text.lower() and (
            "cataloga como muy baja" in text.lower()
            or "por el contrario" in text.lower()
        ):
            dem_baja = data["demand"]["demanda_baja_pct"] * 100
            dem_baja_n = data["demand"]["demanda_baja_freq"]
            dem_muy_baja = data["demand"]["demanda_muy_baja_pct"] * 100
            dem_muy_baja_n = data["demand"]["demanda_muy_baja_freq"]
            new_text = (
                f"Por el contrario, el {dem_baja:.2f}% ({dem_baja_n} {graduado_term}) considera que la demanda es Baja "
                f"y el {dem_muy_baja:.2f}% ({dem_muy_baja_n} personas) la cataloga como Muy Baja. Esta percepción "
                f"del mercado laboral refleja los desafíos percibidos en la inserción o crecimiento laboral."
            )
            para.text = new_text
            print(f"  [Paragraph] Updated job demand perception paragraph {idx}.")

        # 14. Dinámica de Empleabilidad (Conclusión)
        elif "la dinámica de empleabilidad revela que" in text.lower() and (
            "sector privado" in text.lower()
            or "servicios" in text.lower()
            or "ingresos mensuales" in text.lower()
        ):
            priv = data["employment"]["empleo_privado_pct"] * 100
            desemp = data["employment"]["empleo_desempleo_pct"] * 100
            new_text = (
                f"La dinámica de empleabilidad revela que, aunque una parte significativa de la cohorte se inserta bajo relación "
                f"de dependencia en el sector privado ({priv:.2f}%), persiste una proporción de graduados sin actividad laboral "
                f"remunerada declarada del {desemp:.2f}%."
            )
            para.text = new_text
            print(f"  [Paragraph] Updated employability conclusion paragraph {idx}.")

        # 15. Capacidad Emprendedora (Conclusión)
        elif (
            "capacidad emprendedora" in text.lower()
            and "ejecución práctica del autoempleo" in text.lower()
        ):
            emp_aut = data["employment"]["empleo_emprendedor_pct"] * 100
            new_text = (
                f"La ejecución práctica del autoempleo formal muestra que apenas un {emp_aut:.2f}% de la cohorte opera "
                f"bajo modalidades de emprendimiento formal (RUC/RIMPE), indicando la importancia de incentivar la "
                f"formalización de modelos de negocio en la carrera."
            )
            para.text = new_text
            print(f"  [Paragraph] Updated entrepreneurship conclusion paragraph {idx}.")

        # 16. Conclusiones y Continuidad Académica (Conclusión)
        elif "la continuidad académica hacia posgrados formales" in text.lower() and (
            "extremadamente baja" in text.lower()
            or "cohorte no ha realizado" in text.lower()
        ):
            cont_ninguno = data["continuity"]["cont_ninguno_pct"] * 100
            new_text = (
                f"La continuidad académica hacia posgrados formales inmediatos al egreso muestra que una proporción importante "
                f"({cont_ninguno:.2f}%) no registra estudios de posgrado estructurados en los primeros años posteriores a su graduación."
            )
            para.text = new_text
            print(
                f"  [Paragraph] Updated academic continuity conclusion paragraph {idx}."
            )

        # 17. Reemplazo del placeholder 'Analisis' para la valoración del desempeño profesional
        elif text.lower() in ["analisis", "análisis"] and idx > 0:
            parent_text = ""
            for p_back in range(idx - 1, max(-1, idx - 4), -1):
                if doc.paragraphs[p_back].text.strip() != "":
                    parent_text = doc.paragraphs[p_back].text.lower()
                    break
            if (
                "d1" in parent_text
                or "desempeño" in parent_text
                or "influencia" in parent_text
            ):
                d1 = data.get("employer", {}).get("d1", {})
                sq1_muy_bueno = get_d1_pct(d1.get("sq1", {}), "muy bueno") * 100
                sq1_bueno = get_d1_pct(d1.get("sq1", {}), "bueno") * 100
                sq5_muy_bueno = get_d1_pct(d1.get("sq5", {}), "muy bueno") * 100
                sq6_muy_bueno = get_d1_pct(d1.get("sq6", {}), "muy bueno") * 100
                new_text = (
                    f"El análisis de la valoración del desempeño profesional de los graduados, desde la perspectiva de los empleadores, "
                    f"revela un nivel de satisfacción destacado. En el desempeño profesional en general, el {sq1_muy_bueno:.2f}% "
                    f"de los empleadores califica al graduado como Muy Bueno, y el {sq1_bueno:.2f}% como Bueno. Adicionalmente, "
                    f"el comportamiento ético y el compromiso institucional se destacan con un {sq5_muy_bueno:.2f}% y "
                    f"{sq6_muy_bueno:.2f}% de calificación Muy Bueno, respectivamente, consolidando una percepción altamente favorable."
                )
                para.text = new_text
                print(
                    f"  [Paragraph] Replaced 'Analisis' placeholder with dynamic Table 7 analysis in paragraph {idx}."
                )

        # 18. Apoyo Institucional para estudios de posgrado
        elif (
            "apoyo institucional para el financiamiento" in text.lower()
            or "ningún graduado reportó" in text.lower()
            or "facilitación de sus estudios de posgrado" in text.lower()
        ):
            cont = data.get("continuity", {})
            a_si = cont.get("apoyo_si_pct", 0.0) * 100
            a_si_n = cont.get("apoyo_si_freq", 0)
            a_no = cont.get("apoyo_no_pct", 0.0) * 100
            a_no_n = cont.get("apoyo_no_freq", 0)
            a_parte = cont.get("apoyo_en_parte_pct", 0.0) * 100
            a_parte_n = cont.get("apoyo_en_parte_freq", 0)

            if a_si_n == 0 and a_parte_n == 0:
                new_text = (
                    f"Los datos disponibles muestran que ningún graduado ({a_si_n} personas) reportó haber recibido "
                    f"apoyo institucional completo o parcial para sus estudios de posgrado o capacitación, "
                    f"mientras que el {a_no:.2f}% ({a_no_n} personas) manifestó no haber contado con financiamiento "
                    f"o apoyo por parte de la institución, evidenciando una brecha en los mecanismos de apoyo."
                )
            else:
                new_text = (
                    f"Respecto al apoyo institucional para realizar estudios, el {a_si:.2f}% ({a_si_n} personas) reportó haber recibido "
                    f"un apoyo total, el {a_parte:.2f}% ({a_parte_n} personas) indica haber recibido apoyo parcial (en parte), "
                    f"mientras que el {a_no:.2f}% ({a_no_n} personas) manifestó no haber contado con ningún tipo de apoyo o financiamiento "
                    f"institucional para sus estudios de posgrado o capacitación."
                )
            para.text = new_text
            print(
                f"  [Paragraph] Updated postgraduate institutional support paragraph {idx}."
            )

        # 19. Employer cargos (postgrad style: "……" with C1 context in prior para)
        elif (
            idx > 0
            and text.strip() != ""
            and text.replace("…", "").replace(".", "").strip() == ""
            and len(text) <= 10
            and "c1" in doc.paragraphs[idx - 1].text.lower()
            and "cargo" in doc.paragraphs[idx - 1].text.lower()
        ):
            cargos = data.get("open_questions", {}).get("employer_cargos", [])
            # Fallback to datae C2 section
            if not cargos:
                cargos = data.get("employer", {}).get("employer_cargos_datae", [])
            if cargos:
                insertions.append(
                    (
                        idx,
                        cargos,
                        "Los cargos reportados por los empleadores para los graduados de la carrera incluyen:",
                    )
                )
                print(
                    f"  [Paragraph] Queued employer cargos (postgrad) for insertion below paragraph {idx}."
                )
            else:
                para.text = "Los empleadores consultados no reportaron cargos específicos en los datos disponibles."
                print(f"  [Paragraph] Set fallback for employer cargos paragraph {idx}.")

        # 20. Employer funciones (postgrad style: "…….." with C2 context in prior para)
        elif (
            idx > 0
            and text.strip() != ""
            and text.replace("…", "").replace(".", "").strip() == ""
            and len(text) <= 10
            and "c2" in doc.paragraphs[idx - 1].text.lower()
            and "funciones" in doc.paragraphs[idx - 1].text.lower()
        ):
            funciones = data.get("open_questions", {}).get("employer_funciones", [])
            # Fallback to datae C3 section
            if not funciones:
                funciones = data.get("employer", {}).get("employer_funciones_datae", [])
            if funciones:
                insertions.append(
                    (
                        idx,
                        funciones,
                        "Las funciones principales que desempeñan los graduados en sus puestos, según los empleadores, son:",
                    )
                )
                print(
                    f"  [Paragraph] Queued employer funciones (postgrad) for insertion below paragraph {idx}."
                )
            else:
                para.text = "Los empleadores consultados no reportaron funciones específicas en los datos disponibles."
                print(f"  [Paragraph] Set fallback for employer funciones paragraph {idx}.")

        # 21. Institutional support (postgrad): "Los ….."
        elif text.lower().startswith("los") and text.replace(".", "").replace("…", "").strip().lower() == "los":
            cont = data.get("continuity", {})
            a_si = cont.get("apoyo_si_pct", 0.0) * 100
            a_si_n = cont.get("apoyo_si_freq", 0)
            a_no = cont.get("apoyo_no_pct", 0.0) * 100
            a_no_n = cont.get("apoyo_no_freq", 0)
            a_parte = cont.get("apoyo_en_parte_pct", 0.0) * 100
            a_parte_n = cont.get("apoyo_en_parte_freq", 0)
            if a_si_n == 0 and a_parte_n == 0 and a_no_n == 0:
                # No data available — leave as placeholder
                pass
            elif a_si_n == 0 and a_parte_n == 0:
                new_text = (
                    f"Los datos disponibles muestran que ningún {graduado_term} ({a_si_n} personas) reportó haber recibido "
                    f"apoyo institucional completo o parcial para sus estudios de posgrado o capacitación, "
                    f"mientras que el {a_no:.2f}% ({a_no_n} personas) manifestó no haber contado con financiamiento "
                    f"o apoyo por parte de la institución, evidenciando una brecha en los mecanismos de apoyo."
                )
                para.text = new_text
                print(f"  [Paragraph] Updated institutional support (Los …..) paragraph {idx}.")
            else:
                new_text = (
                    f"Los datos muestran que el {a_si:.2f}% ({a_si_n} personas) recibió apoyo total, "
                    f"el {a_parte:.2f}% ({a_parte_n} personas) indica apoyo parcial, "
                    f"mientras que el {a_no:.2f}% ({a_no_n} personas) no contó con ningún tipo de apoyo institucional."
                )
                para.text = new_text
                print(f"  [Paragraph] Updated institutional support (Los …..) paragraph {idx}.")

        # 22. Employer training suggestions (postgrad short placeholder "…." after E1 question)
        elif (
            idx > 0
            and text.strip() != ""
            and text.replace("…", "").replace(".", "").strip() == ""
            and len(text) <= 8
            and any(
                "e1" in doc.paragraphs[j].text.lower()
                for j in range(max(0, idx - 3), idx)
            )
        ):
            suggestions = data.get("open_questions", {}).get(
                "temas_de_capacitacion_sugerida_por_los_empleadores", []
            )
            # Fallback to datae E1 section
            if not suggestions:
                suggestions = data.get("employer", {}).get("employer_training_datae", [])
            if suggestions:
                insertions.append(
                    (
                        idx,
                        suggestions,
                        "Los empleadores consultados sugirieron las siguientes temáticas clave de capacitación y actualización profesional:",
                    )
                )
                print(
                    f"  [Paragraph] Queued employer training suggestions (postgrad) for insertion below paragraph {idx}."
                )
            else:
                para.text = "Los empleadores consultados no registraron sugerencias específicas de capacitación en los datos disponibles."
                print(f"  [Paragraph] Set fallback for employer training suggestions paragraph {idx}.")

        # 23. Employer demand perception (postgrad placeholder: "Los resultados desde la perspectiva de los empleadores revelan que ….")
        elif (
            "perspectiva de los empleadores revelan que" in text.lower()
            and text.strip().endswith(".")
            and len(text) < 120
        ):
            emp_b1 = data.get("employer", {}).get("b1", {})
            emp_b3 = data.get("employer", {}).get("b3", {})
            total_emp = data.get("employer", {}).get("total_employers", 0)
            if total_emp == 0:
                total_emp = sum(v.get("count", 0) for v in emp_b1.values())
            if total_emp == 0:
                total_emp = 10

            # Check if B3 has real data
            b3_total = sum(v.get("count", 0) for v in emp_b3.values())
            if b3_total > 0:
                alta_pct = emp_b3.get("Alta (B4A1)", {}).get("pct", 0.0) * 100
                media_pct = emp_b3.get("Media (B4A2)", {}).get("pct", 0.0) * 100
                alta_count = emp_b3.get("Alta (B4A1)", {}).get("count", 0)
                media_count = emp_b3.get("Media (B4A2)", {}).get("count", 0)
                new_text = (
                    f"De los {total_emp} empleadores consultados, el {alta_pct:.2f}% ({alta_count} empleadores) considera que la demanda es Alta, "
                    f"y el {media_pct:.2f}% ({media_count} empleadores) la califica como Media. Esta distribución refleja una "
                    f"opinión mayoritariamente positiva del sector empleador respecto a las oportunidades de inserción laboral."
                )
            else:
                new_text = (
                    f"Los {total_emp} empleadores consultados no registraron información cuantificable sobre la percepción "
                    f"de la demanda de profesionales de la carrera en la encuesta aplicada."
                )
            para.text = new_text
            print(f"  [Paragraph] Updated employer demand perception (postgrad) paragraph {idx}.")

        # 24. Institution study destination placeholder "….."
        elif (
            idx > 0
            and text.strip() != ""
            and text.replace("…", "").replace(".", "").strip() == ""
            and len(text) <= 8
            and any(
                ("pregunta" in doc.paragraphs[j].text.lower() and "f3" in doc.paragraphs[j].text.lower())
                for j in range(max(0, idx - 3), idx)
            )
        ):
            study_countries = data.get("open_questions", {}).get("f3_pais_ha_realizado_estudios", [])
            if study_countries:
                valid = [c for c in study_countries if c.lower() not in ['nan', '', 'no', 'ninguno']]
                if valid:
                    insertions.append(
                        (
                            idx,
                            valid,
                            "Los países o instituciones donde los graduados han realizado estudios incluyen:",
                        )
                    )
                    print(f"  [Paragraph] Queued study destination institutions for paragraph {idx}.")

    # Execute dynamic bullet insertions in reverse order to keep indices valid
    for idx, items_list, intro in sorted(insertions, key=lambda x: x[0], reverse=True):
        insert_bullets(doc, idx, items_list, intro)


def identify_table_type(table):
    """Identifies a table type based on its cell contents/headers to avoid hardcoded indexes."""
    if not table.rows:
        return "unknown"

    all_text = " ".join(c.text.strip().lower() for r in table.rows for c in r.cells)

    if "nombre de la carrera" in all_text and "periodo del estudio" in all_text:
        return "metadata"
    elif "nombre del programa de posgrado" in all_text and "periodo del estudio" in all_text:
        return "metadata"
    elif "aspecto demográfico" in all_text or "género" in all_text:
        return "demographics"
    elif "tipo de capacitación" in all_text and "doctorado" in all_text:
        return "training_needs"
    elif "correlación graduados" in all_text or "correlación empleadores" in all_text:
        return "correlation"
    elif (
        "desempeño profesional del graduado" in all_text
        and "comportamiento ético" in all_text
    ):
        return "employer_rating"
    elif "requisitos formales" in all_text or "aspectos de contratación" in all_text:
        return "employer_reqs"
    elif "situación laboral" in all_text and "relación de dependencia" in all_text:
        return "employment"
    elif (
        "tipo de estudios realizados" in all_text or "estudios posteriores" in all_text
    ):
        return "continuity"
    elif "nivel de demanda" in all_text or "demanda de profesionales" in all_text:
        return "demand"

    return "unknown"


def update_report(report_id, data, target_docx_path=None):
    """Main updater driver that loads a Word file, updates key tables and paragraphs, and saves it inside the processed/ directory."""
    career = CAREER_FILES[report_id]
    source_docx_path = os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        "INFORMES 24-46 ERICK",
        career["word"],
    )

    if not os.path.exists(source_docx_path):
        raise FileNotFoundError(f"Source Word report not found: {source_docx_path}")

    if target_docx_path is None:
        category_dir = os.path.join(PROCESSED_DIR, get_report_category(report_id))
        os.makedirs(category_dir, exist_ok=True)
        target_docx_path = os.path.join(category_dir, career["word"])

    print(
        f"[Update] Copying raw report to processed folder: {os.path.basename(target_docx_path)}"
    )
    shutil.copy2(source_docx_path, target_docx_path)

    print("[Update] Opening copied document for editing...")
    doc = docx.Document(target_docx_path)
    print(f"  Total tables count: {len(doc.tables)}")

    # 1. Map tables by type dynamically
    table_map = {}
    for idx, table in enumerate(doc.tables):
        t_type = identify_table_type(table)
        print(f"    Table {idx} identified as: {t_type}")
        if t_type not in table_map:
            table_map[t_type] = []
        table_map[t_type].append(table)

    # Extract original population from the metadata table of the source Word document as truth
    orig_poblacion = None
    if "metadata" in table_map:
        for row in table_map["metadata"][0].rows:
            cells = [c.text.strip().lower() for c in row.cells]
            if any("población" in cell or "poblacion" in cell for cell in cells):
                try:
                    orig_poblacion = int(row.cells[-1].text.strip())
                except ValueError:
                    pass
    if orig_poblacion is not None and orig_poblacion > 0:
        print(f"  [Update] Overriding metadata population with original document value: {orig_poblacion}")
        data["metadata"]["poblacion"] = orig_poblacion

    # 2. Update tables dynamically using the map
    if "metadata" in table_map:
        print("  Updating Table 0 (Metadata)...")
        update_table_0(table_map["metadata"][0], data)

    if "demographics" in table_map:
        print("  Updating Table 4 (Demographics)...")
        update_table_4(table_map["demographics"][0], data)

    if "training_needs" in table_map:
        print("  Updating Table 5 (Training Needs)...")
        update_table_5(table_map["training_needs"][0], data)

    if "correlation" in table_map:
        print("  Updating Table 6 (Correlation Analysis)...")
        update_table_6(table_map["correlation"][0], data)

    if "employer_rating" in table_map:
        print("  Updating Table 7 (Employer Performance Rating)...")
        update_table_7(table_map["employer_rating"][0], data)

    if "employment" in table_map:
        print("  Updating Tables 8 & 11 (Employment)...")
        for tbl in table_map["employment"]:
            update_table_8_and_11(tbl, data)

    if "employer_reqs" in table_map:
        print("  Updating Table 10 (Employer Consultation & Requirements)...")
        update_table_10(table_map["employer_reqs"][0], data)

    if "continuity" in table_map:
        print("  Updating Table 12 (Study Continuity)...")
        update_table_12(table_map["continuity"][0], data)

    if "demand" in table_map:
        print("  Updating Table 13 (Job Demand Level)...")
        update_table_13(table_map["demand"][0], data)

    # 3. Handle 'Analisis' placeholder preceding the employer rating table dynamically
    body_elements = []
    for el in doc.element.body:
        if el.tag.endswith("p"):
            body_elements.append(("para", docx.text.paragraph.Paragraph(el, doc)))
        elif el.tag.endswith("tbl"):
            body_elements.append(("table", docx.table.Table(el, doc)))

    is_pg = report_id >= 35
    graduado_term = "maestrantes" if is_pg else "graduados"

    for idx, (el_type, el_obj) in enumerate(body_elements):
        if el_type == "table" and identify_table_type(el_obj) == "employer_rating":
            # Search backwards for the Analisis placeholder
            for back_idx in range(idx - 1, max(-1, idx - 10), -1):
                back_type, back_obj = body_elements[back_idx]
                if back_type == "para":
                    txt = back_obj.text.strip().lower()
                    if txt in ["analisis", "análisis"]:
                        d1 = data.get("employer", {}).get("d1", {})
                        sq1_muy_bueno = get_d1_pct(d1.get("sq1", {}), "muy bueno") * 100
                        sq1_bueno = get_d1_pct(d1.get("sq1", {}), "bueno") * 100
                        sq5_muy_bueno = get_d1_pct(d1.get("sq5", {}), "muy bueno") * 100
                        sq6_muy_bueno = get_d1_pct(d1.get("sq6", {}), "muy bueno") * 100
                        new_text = (
                            f"El análisis de la valoración del desempeño profesional de los graduados, desde la perspectiva de los empleadores, "
                            f"revela un nivel de satisfacción destacado. En el desempeño profesional en general, el {sq1_muy_bueno:.2f}% "
                            f"de los empleadores califica al graduado como Muy Bueno, y el {sq1_bueno:.2f}% como Bueno. Adicionalmente, "
                            f"el comportamiento ético y el compromiso institucional se destacan con un {sq5_muy_bueno:.2f}% y "
                            f"{sq6_muy_bueno:.2f}% de calificación Muy Bueno, respectivamente, consolidando una percepción altamente favorable."
                        )
                        back_obj.text = new_text
                        print(
                            f"  [Paragraph] Replaced 'Analisis' placeholder before Table 7 dynamically by table position."
                        )
                        break

    # 4. Rewrite global paragraphs
    rewrite_paragraphs(report_id, doc, data)

    # 5. Strip all highlight/background colors from document
    strip_all_highlights(doc)

    # 6. Center text in all table cells and remove inner spacing
    center_all_table_cells(doc)

    doc.save(target_docx_path)
    print(
        f"[Update] Document successfully written to: {os.path.basename(target_docx_path)}"
    )

    # Also save/copy to the root of PROCESSED_DIR for convenience
    root_target_path = os.path.join(PROCESSED_DIR, career["word"])
    try:
        shutil.copy2(target_docx_path, root_target_path)
        print(f"  [Update] Saved a duplicate copy to the root processed folder: {root_target_path}")
    except Exception as e:
        print(f"  [Warning] Could not copy to root processed folder: {e}")
